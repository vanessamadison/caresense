"""Secure document parsing for multiple formats."""

from __future__ import annotations

import hashlib
import mimetypes
from functools import lru_cache
from pathlib import Path
from typing import Dict, Optional

import docx
from email_reply_parser import EmailReplyParser
from pypdf import PdfReader

from caresense.parsers.sanitizer import get_sanitizer
from caresense.utils.logging import get_logger

log = get_logger(__name__)

# Security: File size limits (in bytes)
MAX_PDF_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_EMAIL_SIZE = 5 * 1024 * 1024  # 5 MB

# Security: Page/section limits
MAX_PDF_PAGES = 50
MAX_DOCX_PARAGRAPHS = 1000
MAX_EMAIL_PARTS = 20

# Allowed MIME types
ALLOWED_MIMETYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "message/rfc822",
}


class DocumentParser:
    """
    Secure multi-format document parser.

    Supports: PDF, DOCX, plain text, email

    Security features:
    - File size validation
    - MIME type validation
    - Content sanitization
    - Resource limits (pages, paragraphs)
    - Error handling
    - Audit logging
    """

    def __init__(self) -> None:
        self._sanitizer = get_sanitizer()
        self._logger = get_logger(__name__)

    def parse(self, file_path: str | Path, source_type: Optional[str] = None) -> Dict[str, any]:
        """
        Parse document and extract text securely.

        Args:
            file_path: Path to document file
            source_type: Optional explicit source type (pdf, docx, email, text)

        Returns:
            Dict containing:
                - text: Extracted and sanitized text
                - metadata: File metadata (type, size, pages/sections)
                - file_hash: SHA256 hash for audit trail
                - pii_detected: PII detection results

        Raises:
            ValueError: If file invalid, too large, or wrong type
            IOError: If file cannot be read

        Security:
            - Validates file exists and is readable
            - Checks file size before loading
            - Validates MIME type
            - Sanitizes all extracted text
            - Detects PII
            - Logs parsing operations
        """
        file_path = Path(file_path)

        # Security: Validate file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.is_file():
            raise ValueError(f"Not a file: {file_path}")

        # Security: Check file size
        file_size = file_path.stat().st_size
        if file_size == 0:
            raise ValueError("Empty file")

        # Compute file hash for audit
        file_hash = self._compute_file_hash(file_path)

        # Determine file type
        if source_type is None:
            mime_type, _ = mimetypes.guess_type(str(file_path))
            source_type = self._mime_to_source_type(mime_type)
        else:
            source_type = source_type.lower()

        # Security: Validate file size based on type
        self._validate_file_size(file_size, source_type)

        self._logger.info(
            "parsing_document",
            file_name=file_path.name,
            file_size=file_size,
            source_type=source_type,
            file_hash=file_hash[:16],
        )

        # Parse based on type
        try:
            if source_type == "pdf":
                text, metadata = self._parse_pdf(file_path)
            elif source_type == "docx":
                text, metadata = self._parse_docx(file_path)
            elif source_type == "email":
                text, metadata = self._parse_email(file_path)
            elif source_type == "text":
                text, metadata = self._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported source type: {source_type}")

            # Security: Sanitize extracted text
            text = self._sanitizer.sanitize(text)

            # Security: Validate it looks like medical content
            if not self._sanitizer.validate_medical_text(text):
                self._logger.warning("suspicious_content_detected", file_hash=file_hash[:16])

            # Security: Detect PII
            pii_detected = self._sanitizer.detect_pii(text)

            result = {
                "text": text,
                "metadata": {
                    **metadata,
                    "file_name": file_path.name,
                    "file_size": file_size,
                    "source_type": source_type,
                },
                "file_hash": file_hash,
                "pii_detected": pii_detected,
            }

            self._logger.info(
                "document_parsed_successfully",
                file_hash=file_hash[:16],
                text_length=len(text),
                pii_found=any(pii_detected.values()),
            )

            return result

        except Exception as e:
            self._logger.error(
                "document_parsing_failed",
                error=str(e),
                file_name=file_path.name,
                source_type=source_type,
            )
            raise

    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file for audit trail."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _mime_to_source_type(self, mime_type: Optional[str]) -> str:
        """Convert MIME type to source type."""
        if mime_type == "application/pdf":
            return "pdf"
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        elif mime_type and mime_type.startswith("text/"):
            return "text"
        elif mime_type == "message/rfc822":
            return "email"
        else:
            raise ValueError(f"Unsupported MIME type: {mime_type}")

    def _validate_file_size(self, size: int, source_type: str) -> None:
        """Validate file size based on type."""
        limits = {
            "pdf": MAX_PDF_SIZE,
            "docx": MAX_DOCX_SIZE,
            "email": MAX_EMAIL_SIZE,
            "text": MAX_EMAIL_SIZE,
        }

        max_size = limits.get(source_type, MAX_EMAIL_SIZE)
        if size > max_size:
            raise ValueError(
                f"File size {size} bytes exceeds maximum {max_size} bytes for {source_type}"
            )

    def _parse_pdf(self, file_path: Path) -> tuple[str, Dict]:
        """Parse PDF file securely."""
        try:
            reader = PdfReader(str(file_path))

            # Security: Limit pages
            num_pages = len(reader.pages)
            if num_pages > MAX_PDF_PAGES:
                self._logger.warning("pdf_too_many_pages", num_pages=num_pages)
                raise ValueError(f"PDF has too many pages: {num_pages} > {MAX_PDF_PAGES}")

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages[:MAX_PDF_PAGES]):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            text = "\n\n".join(text_parts)

            metadata = {
                "num_pages": num_pages,
                "pdf_version": reader.pdf_header,
            }

            return text, metadata

        except Exception as e:
            self._logger.error("pdf_parsing_failed", error=str(e))
            raise ValueError(f"Failed to parse PDF: {e}")

    def _parse_docx(self, file_path: Path) -> tuple[str, Dict]:
        """Parse DOCX file securely."""
        try:
            doc = docx.Document(str(file_path))

            # Security: Limit paragraphs
            num_paragraphs = len(doc.paragraphs)
            if num_paragraphs > MAX_DOCX_PARAGRAPHS:
                self._logger.warning("docx_too_many_paragraphs", num_paragraphs=num_paragraphs)
                raise ValueError(
                    f"DOCX has too many paragraphs: {num_paragraphs} > {MAX_DOCX_PARAGRAPHS}"
                )

            # Extract text from paragraphs
            text_parts = [para.text for para in doc.paragraphs[:MAX_DOCX_PARAGRAPHS] if para.text.strip()]

            text = "\n\n".join(text_parts)

            metadata = {
                "num_paragraphs": num_paragraphs,
            }

            return text, metadata

        except Exception as e:
            self._logger.error("docx_parsing_failed", error=str(e))
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _parse_email(self, file_path: Path) -> tuple[str, Dict]:
        """Parse email file securely."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                email_content = f.read()

            # Security: Validate size
            if len(email_content) > MAX_EMAIL_SIZE:
                raise ValueError("Email content too large")

            # Use email reply parser to extract main content
            parsed_email = EmailReplyParser.parse_reply(email_content)

            metadata = {
                "email_length": len(email_content),
                "parsed_length": len(parsed_email),
            }

            return parsed_email, metadata

        except Exception as e:
            self._logger.error("email_parsing_failed", error=str(e))
            raise ValueError(f"Failed to parse email: {e}")

    def _parse_text(self, file_path: Path) -> tuple[str, Dict]:
        """Parse plain text file securely."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

            metadata = {
                "text_length": len(text),
            }

            return text, metadata

        except Exception as e:
            self._logger.error("text_parsing_failed", error=str(e))
            raise ValueError(f"Failed to parse text: {e}")


@lru_cache
def get_document_parser() -> DocumentParser:
    """Return singleton document parser instance."""
    return DocumentParser()
